# -*- coding: utf-8 -*-
"""生成课程设计报告 Word 文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = r"C:\Users\张昊\Doubao\chats\2026-08-29\new-chat\zh111\工业设备轴承故障智能监测与诊断系统_课程设计报告.docx"

doc = Document()

# ========== 页面设置 ==========
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ========== 样式设置 ==========
def set_font(run, name_cn, name_en, size, bold=False, color=None):
    run.font.name = name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name_cn)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

# 正文样式
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.first_line_indent = Pt(24)  # 2字符
normal.paragraph_format.space_after = Pt(0)

# 标题样式
for level, (size, space_before, space_after) in [
    (1, (16, 18, 12)),
    (2, (14, 14, 8)),
    (3, (12, 10, 6)),
]:
    style = doc.styles[f'Heading {level}']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_before = Pt(space_before)
    style.paragraph_format.space_after = Pt(space_after)
    style.paragraph_format.first_line_indent = Pt(0)
    style.paragraph_format.line_spacing = 1.5

def add_para(text, style='Normal', align=None, indent=True):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    if not indent:
        p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    return p

def add_heading(text, level):
    h = doc.add_heading(text, level=level)
    return h

# ========== 封面 ==========
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
run = p.add_run('制造智能技术课程设计报告')
set_font(run, '黑体', 'Times New Roman', 22, bold=True)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
run = p.add_run('工业设备轴承故障智能监测与诊断系统')
set_font(run, '黑体', 'Times New Roman', 26, bold=True)

for _ in range(6):
    doc.add_paragraph()

info_lines = [
    ('学    院', '智能制造学院'),
    ('专    业', '智能制造工程'),
    ('班    级', '智造24-2-25'),
    ('姓    名', '张昊'),
    ('完成日期', '2026 年 9 月'),
]
for label, value in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(f'{label}：{value}')
    set_font(run, '宋体', 'Times New Roman', 14)

doc.add_page_break()

# ========== 摘要 ==========
add_heading('摘  要', level=1)

abstract = (
    '滚动轴承是工业旋转机械中应用最广泛的核心部件之一，其运行状态直接影响设备的可靠性与安全性。'
    '传统的人工巡检和定期维护存在效率低、滞后性强等问题，难以满足现代工业对设备健康管理的需求。'
    '本课程设计基于 CWRU（凯斯西储大学）公开轴承振动数据集，构建了一个 B/S 架构的轴承故障智能监测与诊断系统。'
    '系统通过带通滤波、去趋势、滑动窗口和 z-score 标准化等预处理步骤，从振动信号中提取 12 维时域特征和 8 维频域特征，'
    '采用随机森林分类器实现正常、内圈故障、滚动体故障和外圈故障四类状态的自动识别。'
    '为提升诊断稳定性，系统采用多窗口投票策略，对全部滑动窗口的预测概率取均值后判定类别。'
    '在健康管理方面，系统基于 RMS 指标构建健康退化序列，通过多项式回归外推实现趋势预测与阈值预警。'
    '同时，系统构建了包含四类故障成因、处理建议和检查清单的知识库，结合峭度等特征进行严重度分级，输出可解释的运维决策。'
    'Web 系统采用 Flask 后端提供 8 个 REST API，前端使用 ECharts 实现波形、频谱和健康趋势三图联动，数据持久化于 SQLite 数据库。'
    '实验结果表明，随机森林模型在随机划分下准确率达 0.997，5 折交叉验证准确率为 0.891。'
    '本项目的核心方法论贡献在于发现并正确处理了滑动窗口重叠导致的数据泄漏问题：'
    '采用按文件分组的 GroupShuffleSplit 评估方法后，真实泛化准确率为 0.633，揭示了随机划分评估的严重虚高问题。'
    '系统 API 测试 12/12 全部通过，浏览器端四类故障诊断验证正确，具备完整的工程可用性。'
)
add_para(abstract)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(0)
run = p.add_run('关键词：')
set_font(run, '黑体', 'Times New Roman', 12, bold=True)
run = p.add_run('轴承故障诊断；随机森林；特征提取；数据泄漏；健康预测；知识推理；Flask')
set_font(run, '宋体', 'Times New Roman', 12)

doc.add_page_break()

# ========== 目录 ==========
add_heading('目  录', level=1)
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(0)
run = p.add_run()
fldChar = OxmlElement('w:fldChar')
fldChar.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:t')
fldChar3.text = "右键更新域以生成目录"
fldChar4 = OxmlElement('w:fldChar')
fldChar4.set(qn('w:fldCharType'), 'end')
run._r.append(fldChar)
run._r.append(instrText)
run._r.append(fldChar2)
run._r.append(fldChar3)
run._r.append(fldChar4)

doc.add_page_break()

# ========== 第1章 绪论 ==========
add_heading('第1章 绪论', level=1)

add_heading('1.1 研究背景与意义', level=2)
add_para(
    '滚动轴承作为旋转机械中最关键的基础部件之一，广泛应用于电动机、风机、泵、机床、齿轮箱和轨道交通等工业设备中。'
    '据统计，旋转机械的故障中约有 30% 与轴承直接相关。轴承一旦发生故障，可能导致设备振动加剧、精度下降，'
    '严重时甚至引发非计划停机和重大经济损失。因此，对轴承运行状态进行实时监测与智能诊断，'
    '实现故障的早期发现和精准定位，对于保障工业设备安全运行、降低维护成本具有重要的工程价值。'
)
add_para(
    '传统的轴承故障诊断主要依赖人工巡检和定期维护。人工巡检依靠运维人员的经验判断，效率低且主观性强；'
    '定期维护则可能造成过度维修或漏检，无法准确反映设备的实际健康状态。随着传感器技术和机器学习的发展，'
    '基于振动信号分析的智能故障诊断方法逐渐成为研究热点。通过采集轴承运行时的振动加速度信号，'
    '提取时域和频域特征，利用分类算法自动识别故障类型，可以大幅提升诊断的效率和准确性。'
)
add_para(
    '本课程设计旨在构建一个集故障诊断、健康趋势预测和知识推理于一体的轴承故障智能监测与诊断系统。'
    '系统不仅能够自动识别轴承故障类型，还能预测健康退化趋势并给出可解释的运维建议，'
    '为设备运维人员提供科学的决策支持。'
)

add_heading('1.2 国内外研究现状', level=2)
add_para(
    '轴承故障诊断技术经历了从传统信号处理到智能诊断的发展历程。早期方法主要基于时域统计指标（如 RMS、峭度）'
    '和频域分析（如 FFT、包络解调），通过人工设定阈值进行故障判定。这类方法原理简单，但对复杂故障的识别能力有限。'
)
add_para(
    '近年来，机器学习方法在轴承故障诊断中得到广泛应用。支持向量机（SVM）、K 近邻（KNN）、随机森林等传统机器学习算法'
    '在手工特征的基础上取得了较好的分类效果。深度学习方法如一维卷积神经网络（1D-CNN）、循环神经网络（LSTM）'
    '和 Transformer 等，能够自动从原始信号中学习特征，在大规模数据集上表现优异。'
    '然而，深度学习方法通常需要大量标注数据和计算资源，且模型可解释性较差，在工业落地中存在一定挑战。'
)
add_para(
    '在评估方法方面，许多研究采用随机划分训练集和测试集的方式评估模型性能。'
    '但由于滑动窗口切片导致相邻样本高度相关，随机划分会产生严重的数据泄漏问题，使评估结果虚高。'
    '近年来，越来越多的研究者开始关注按设备或文件分组的评估方法，以更真实地反映模型的泛化能力。'
)

add_heading('1.3 本文主要工作', level=2)
add_para(
    '本课程设计的主要工作包括以下几个方面：'
)
add_para(
    '（1）数据处理与特征提取：基于 CWRU 公开数据集，完成带通滤波、去趋势、滑动窗口切片和 z-score 标准化等预处理，'
    '提取 12 维时域特征和 8 维频域特征，共生成 6389 个训练样本。'
)
add_para(
    '（2）故障分类算法：对比 KNN、SVM 和随机森林三种分类器，选择随机森林作为最终模型，'
    '采用多窗口投票策略提升诊断稳定性。重点研究了滑动窗口重叠导致的数据泄漏问题，'
    '对比了随机划分和按文件分组两种评估方法的差异。'
)
add_para(
    '（3）健康预测与知识推理：基于 RMS 指标构建健康退化序列，通过多项式回归外推实现趋势预测；'
    '构建四类故障知识库，结合峭度特征进行严重度分级，输出可解释的运维建议。'
)
add_para(
    '（4）Web 系统开发：采用 Flask 后端和 ECharts 前端，实现数据加载、故障诊断、健康预测、知识推理和记录查询等功能，'
    '使用 SQLite 进行数据持久化，完成系统联调和功能测试。'
)

# ========== 第2章 系统总体设计 ==========
add_heading('第2章 系统总体设计', level=1)

add_heading('2.1 需求分析', level=2)
add_para(
    '本系统面向工业设备轴承故障监测场景，主要功能需求包括：'
    '（1）支持 CWRU 格式 .mat 振动数据文件的加载与可视化展示；'
    '（2）对振动信号进行自动故障诊断，识别正常、内圈故障、滚动体故障和外圈故障四类状态；'
    '（3）对轴承健康状态进行趋势预测，当预测值超过阈值时触发预警；'
    '（4）根据诊断结果输出故障成因、处理建议和检查清单等可解释的运维决策；'
    '（5）支持诊断记录和告警记录的持久化存储与历史查询。'
)
add_para(
    '非功能需求方面，系统应具备良好的响应速度和稳定性，前端界面应简洁直观，'
    '支持离线演示（前端资源本地化），后端 API 应具备完整的错误处理能力。'
)

add_heading('2.2 系统架构', level=2)
add_para(
    '系统采用 B/S（浏览器/服务器）三层架构，自上而下分为表现层、应用层、算法层和数据层。'
)
add_para(
    '表现层运行于浏览器端，使用 HTML、CSS、JavaScript 和 ECharts 实现数据可视化，'
    '包括时域波形图、频谱分析图和健康趋势图的三图联动展示，以及诊断结果和运维建议的呈现。'
)
add_para(
    '应用层基于 Flask Web 框架，提供 8 个 REST API 接口，负责接收前端请求、调用算法模块、'
    '操作数据库并返回 JSON 格式的响应数据。'
)
add_para(
    '算法层以 Python 模块形式独立实现，包含信号预处理、特征提取、故障分类、健康预测和知识推理五个子模块，'
    '与 Web 层解耦，便于独立测试和维护。'
)
add_para(
    '数据层包含 CWRU 原始 .mat 振动数据文件和 SQLite 数据库。SQLite 数据库包含传感器数据、故障样本、'
    '知识库和告警记录四张表，支持诊断结果的持久化存储和历史追溯。'
)

add_heading('2.3 技术选型', level=2)
add_para(
    '后端选择 Flask 框架，因其轻量灵活、学习成本低，适合课程设计场景。'
    '数据库选择 SQLite，零配置、单文件存储，便于项目提交和答辩演示。'
    '前端可视化选择 ECharts，其图表类型丰富、交互能力强，且支持本地部署避免 CDN 依赖。'
    '机器学习库使用 scikit-learn，提供成熟的随机森林、SVM、KNN 等分类器实现。'
    '信号处理使用 SciPy 和 NumPy，特征计算高效可靠。'
)

# ========== 第3章 数据处理与特征提取 ==========
add_heading('第3章 数据处理与特征提取', level=1)

add_heading('3.1 数据集介绍', level=2)
add_para(
    '本项目使用 CWRU（Case Western Reserve University）轴承数据集，这是轴承故障诊断领域最常用的公开基准数据集之一。'
    '该数据集在 2 马力电动机上，通过电火花加工在轴承内圈、滚动体和外圈注入不同尺寸（0.007~0.028 英寸）的故障，'
    '在 0~3 马力负载条件下采集驱动端和风扇端的振动加速度信号，采样率为 12 kHz 和 48 kHz。'
)
add_para(
    '本项目选取了 20 个 .mat 数据文件，涵盖四类状态：正常（3 个文件）、内圈故障（9 个文件）、'
    '滚动体故障（3 个文件）和外圈故障（5 个文件）。数据目录已加入 .gitignore，避免将大数据文件提交至版本控制系统。'
    '数据集来源链接已在项目 README 中明确标注，遵循公开数据集引用规范。'
)

add_heading('3.2 信号预处理', level=2)
add_para(
    '原始振动信号包含噪声和基线漂移，需要进行预处理以提升后续特征提取和分类的准确性。预处理流程包括以下四个步骤：'
)
add_para(
    '（1）带通滤波：使用 Butterworth 带通滤波器去除信号中的低频漂移和高频噪声，保留轴承故障特征频率所在的频带。'
)
add_para(
    '（2）去趋势：消除信号中的线性趋势项，减少基线漂移对时域特征计算的影响。'
)
add_para(
    '（3）滑动窗口切片：采用窗口长度 1024 点、步长 512 点的滑动窗口对信号进行切片。'
    '步长小于窗口长度意味着相邻窗口存在 50% 的重叠，这可以增加样本数量，但也引入了数据泄漏风险（详见第 4.3 节）。'
)
add_para(
    '（4）z-score 标准化：对每个窗口的信号进行零均值单位方差标准化，消除不同工况下信号幅值差异的影响。'
)
add_para(
    '经过上述预处理，共生成 6389 个训练样本，每个样本包含 1024 个采样点。'
)

add_heading('3.3 特征提取', level=2)
add_para(
    '从每个滑动窗口中提取 20 维特征向量，包括 12 维时域特征和 8 维频域特征。'
)
add_para(
    '时域特征包括：均值、绝对均值、标准差、方差、均方根（RMS）、峰峰值、峰值、峭度（Kurtosis）、'
    '偏度（Skewness）、峰值因子、脉冲因子和裕度因子。其中，峭度是衡量信号冲击特性的重要指标，'
    '正常轴承的峭度接近 3，当出现故障冲击时峭度会显著升高，因此是故障诊断中最敏感的特征之一。'
)
add_para(
    '频域特征通过快速傅里叶变换（FFT）将信号转换到频域后计算，包括：频谱幅值均值、频谱幅值标准差、'
    '频谱质心、频谱 RMS、主频幅值、主频位置、频谱峭度和频谱偏度。频域特征能够反映故障特征频率的分布情况，'
    '有助于区分内圈、外圈和滚动体故障。'
)
add_para(
    '所有特征在输入分类器前进行标准化处理，确保不同量纲的特征对模型的贡献均衡。'
)

# ========== 第4章 故障诊断算法 ==========
add_heading('第4章 故障诊断算法', level=1)

add_heading('4.1 分类模型选择', level=2)
add_para(
    '本项目对比了三种经典的机器学习分类器：K 近邻（KNN）、支持向量机（SVM，RBF 核）和随机森林（Random Forest）。'
    'KNN 基于距离度量进行分类，原理简单但对高维数据和噪声敏感；SVM 通过核函数将数据映射到高维空间寻找最优分类超平面，'
    '在小样本场景下表现较好；随机森林通过集成多棵决策树进行投票，具有训练速度快、抗过拟合能力强、'
    '可输出特征重要性等优点。'
)
add_para(
    '实验结果表明，随机森林在 5 折交叉验证中取得了最高的准确率（0.891），因此被选为最终分类模型。'
    '随机森林设置 100 棵决策树，通过特征随机子空间和样本自助采样（Bootstrap）降低单棵树的过拟合风险。'
)

add_heading('4.2 实验结果与分析', level=2)
add_para(
    '三种分类器的实验结果如表 4-1 所示。'
)

# 表格
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
table.alignment = WD_ALIGN_PARAGRAPH.CENTER
headers = ['模型', '5 折交叉验证准确率', '随机划分测试准确率']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(h)
    set_font(run, '黑体', 'Times New Roman', 11, bold=True)
data_rows = [
    ['KNN', '0.842', '0.973'],
    ['SVM (RBF)', '0.868', '0.985'],
    ['随机森林', '0.891', '0.997'],
]
for r, row_data in enumerate(data_rows):
    for c, val in enumerate(row_data):
        cell = table.rows[r+1].cells[c]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(val)
        set_font(run, '宋体', 'Times New Roman', 11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
run = p.add_run('表 4-1  三种分类器性能对比')
set_font(run, '宋体', 'Times New Roman', 10.5)

add_para(
    '为提升诊断稳定性，系统采用多窗口投票策略：对信号的全部滑动窗口分别提取特征并预测，'
    '取各窗口 predict_proba 的均值，最大概率对应的类别作为最终诊断结果，最大概率值作为置信度。'
    '相比单窗口预测，多窗口投票能够有效降低个别异常窗口对诊断结果的影响，显著提升稳定性。'
)

add_heading('4.3 数据泄漏问题与正确评估', level=2)
add_para(
    '本项目最重要的方法论发现是关于数据泄漏问题的深入分析。'
    '在最初的实验中，采用随机划分训练集和测试集的方法，随机森林的测试准确率高达 0.997，'
    '这一结果看似非常优秀，但经过深入分析后发现存在严重的数据泄漏问题。'
)
add_para(
    '数据泄漏的根源在于滑动窗口的重叠设置：窗口长度为 1024 点，步长为 512 点，'
    '相邻窗口之间有 50% 的数据重叠，因此来自同一原始文件的窗口之间高度相关。'
    '当采用随机划分时，同一文件的窗口可能同时出现在训练集和测试集中，'
    '模型在训练时已经"见过"测试样本的高度相似数据，导致评估结果虚高。'
)
add_para(
    '为正确评估模型的泛化能力，本项目采用按文件分组的 GroupShuffleSplit 方法：'
    '以原始数据文件为分组单位，确保同一文件的所有窗口不会同时出现在训练集和测试集中。'
    '这种评估方式更接近真实工业场景——模型需要对"未见设备"的数据进行诊断。'
)
add_para(
    '实验结果显示，按文件分组评估的准确率为 0.633，与随机划分的 0.997 相差 36.4 个百分点。'
    '这一巨大差距充分说明，随机划分严重高估了模型的实际泛化能力。'
    '在课程设计和工业部署中，应采用按设备或文件分组的评估方法，以获得真实可靠的性能评估。'
    '这一发现是本项目的核心方法论贡献，也为后续研究提供了重要的评估规范参考。'
)

# ========== 第5章 健康预测与知识推理 ==========
add_heading('第5章 健康预测与知识推理', level=1)

add_heading('5.1 健康趋势预测', level=2)
add_para(
    '除了故障类型识别，系统还提供轴承健康趋势预测功能，帮助运维人员提前发现设备退化趋势。'
    '健康预测的实现步骤如下：'
)
add_para(
    '（1）构建健康序列：对振动信号进行滑动窗口切片，计算每个窗口的 RMS（均方根）值，'
    '形成随时间变化的健康指标序列。RMS 值能够反映振动能量的整体水平，是衡量轴承健康状态的常用指标。'
)
add_para(
    '（2）趋势拟合与外推：使用多项式回归对健康序列进行拟合，并外推未来 30 步的趋势值。'
    '多项式回归能够捕捉健康指标的非线性退化趋势。'
)
add_para(
    '（3）健康度计算：将 RMS 值归一化到 0~100 的健康度范围，100 表示完全健康，0 表示完全失效。'
)
add_para(
    '（4）预警判定：基于健康序列的基线（均值加倍数标准差）动态计算预警阈值，'
    '当预测的末端值超过阈值时，系统判定为预警状态。'
)
add_para(
    '实测结果表明，内圈故障数据的健康度降至 35.3，系统正确触发了重度告警，验证了健康预测模块的有效性。'
)

add_heading('5.2 严重度分级', level=2)
add_para(
    '系统基于峭度（Kurtosis）指标对故障严重程度进行分级。峭度是衡量信号冲击特性的统计量，'
    '正常轴承的峭度约为 3，故障冲击越强峭度值越高。分级规则如下：'
)
add_para(
    '轻度（峭度 < 5）：故障初期，建议纳入例行检修计划；'
    '中度（5 ≤ 峭度 < 8）：故障发展期，建议近期安排检修；'
    '重度（峭度 ≥ 8）：故障严重期，建议立即停机处理。'
)
add_para(
    '严重度分级与故障类型诊断相结合，能够为运维人员提供更精准的决策依据。'
)

add_heading('5.3 知识库与运维建议', level=2)
add_para(
    '为提升系统的可解释性，本项目构建了包含四类故障的知识库，每类故障记录了名称、成因、特征描述、'
    '处理建议和检查清单。知识库存储于 SQLite 数据库的 fault_knowledge 表中。'
)
add_para(
    '以内圈故障为例，知识库记录的成因包括内圈滚道点蚀、剥落和疲劳裂纹，常见于装配过紧、'
    '润滑不良或长期过载；处理建议为早期密切监控、中重度更换轴承；'
    '检查清单包括检查内圈滚道剥落/点蚀、核查装配过盈量与润滑状态、重度时更换轴承。'
)
add_para(
    '诊断完成后，系统根据预测的故障类型查询知识库，结合严重度分级，输出包含故障成因、'
    '处理建议和检查清单的完整运维决策。这种"诊断结果 + 知识推理"的模式使系统不仅能告诉用户"是什么故障"，'
    '还能解释"为什么"和"怎么办"，更贴近工业运维的实际需求。'
)

# ========== 第6章 Web系统实现 ==========
add_heading('第6章 Web 系统实现', level=1)

add_heading('6.1 后端设计', level=2)
add_para(
    '后端基于 Flask 框架开发，共提供 8 个 REST API 接口，覆盖健康检查、数据集列表、信号读取、'
    '故障诊断、健康预测、知识推理、知识库查询和记录查询等功能。各接口的功能如表 6-1 所示。'
)

table2 = doc.add_table(rows=9, cols=3)
table2.style = 'Table Grid'
table2.alignment = WD_ALIGN_PARAGRAPH.CENTER
api_headers = ['接口', '方法', '功能']
for i, h in enumerate(api_headers):
    cell = table2.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(h)
    set_font(run, '黑体', 'Times New Roman', 11, bold=True)
api_data = [
    ['/api/health', 'GET', '服务健康检查'],
    ['/api/datasets', 'GET', '获取数据集文件列表'],
    ['/api/signal/<file>', 'GET', '获取波形与频谱数据'],
    ['/api/diagnose', 'POST', '执行故障诊断'],
    ['/api/predict', 'POST', '执行健康趋势预测'],
    ['/api/reason', 'POST', '知识推理与严重度分级'],
    ['/api/knowledge', 'GET', '查询故障知识库'],
    ['/api/records', 'GET', '查询诊断与告警记录'],
]
for r, row_data in enumerate(api_data):
    for c, val in enumerate(row_data):
        cell = table2.rows[r+1].cells[c]
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(val)
        set_font(run, '宋体', 'Times New Roman', 10.5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
run = p.add_run('表 6-1  REST API 接口一览')
set_font(run, '宋体', 'Times New Roman', 10.5)

add_para(
    '后端的核心设计原则是算法模块解耦：Web 层仅调用 src/ 目录中已实现的算法函数，不重复实现业务逻辑。'
    '诊断接口的处理流程为：接收信号数据 → 预处理 → 特征提取 → 随机森林多窗口投票 → 知识推理 → 写入数据库 → 返回结果。'
    '对于中重度故障，系统自动写入告警记录表。'
)
add_para(
    '针对原始信号数据量大（最长约 24 万点）的问题，后端采用等间隔抽稀策略：'
    '波形数据抽稀至 3000 点，频谱数据先计算全量 FFT 后取 0~6000Hz 范围再抽稀至 1500 点，'
    '有效降低了前端渲染压力。'
)

add_heading('6.2 前端设计', level=2)
add_para(
    '前端采用 HTML + CSS + JavaScript 开发，使用 Bootstrap 进行响应式布局，ECharts 实现数据可视化。'
    'ECharts 库已下载至本地 static/js/ 目录，避免 CDN 不可达导致页面白屏，保障离线演示可用。'
)
add_para(
    '前端主要组件包括：数据选择面板（20 个数据文件下拉框 + 加载/诊断/预测按钮）、'
    'ECharts 波形图（驱动端振动波形，支持缩放）、ECharts 频谱图（0~6000Hz 幅值谱）、'
    'ECharts 趋势图（历史健康指标 + 趋势预测 + 预警阈值三线叠加）、'
    '健康度面板（当前健康度与预警状态实时联动）、诊断结果与运维建议面板、'
    '以及诊断记录和告警记录表格。三个图表通过共享的数据集选择实现联动，'
    '用户选择数据后波形和频谱自动刷新，点击诊断或预测按钮后趋势图和结果面板同步更新。'
)

add_heading('6.3 数据库设计', level=2)
add_para(
    '系统使用 SQLite 数据库进行数据持久化，包含四张表：'
)
add_para(
    '（1）sensor_data 表：存储振动传感器时序数据，字段包括时间戳、设备 ID、通道和振动值。'
)
add_para(
    '（2）fault_sample 表：存储故障样本记录，字段包括采样时间、标签、故障类型、特征向量和置信度。'
)
add_para(
    '（3）fault_knowledge 表：存储故障知识库，字段包括故障类型、名称、成因、特征描述、处理建议和检查清单，'
    '共预置 4 条记录，故障类型字段设为 UNIQUE 约束。'
)
add_para(
    '（4）alarm_record 表：存储告警记录，字段包括时间戳、设备 ID、告警类型和描述。'
)
add_para(
    '选择 SQLite 的原因是课程设计单机部署场景下，SQLite 零配置、单文件、够用且便于提交答辩，'
    '避免了 MySQL 等数据库的服务器部署复杂度。'
)

# ========== 第7章 系统测试与验证 ==========
add_heading('第7章 系统测试与验证', level=1)

add_heading('7.1 API 接口测试', level=2)
add_para(
    '系统编写了自动化接口测试脚本 scripts/test_api.py，对 8 个 API 共设计了 12 个测试用例，'
    '覆盖数据加载、故障诊断、健康预测、知识推理和记录查询等核心功能。测试结果为 12/12 全部通过，'
    '所有接口均能正确返回预期的 JSON 响应，错误处理机制正常工作。'
)

add_heading('7.2 端到端验证', level=2)
add_para(
    '在浏览器端进行了完整的端到端功能测试。测试流程为：启动 Flask 服务 → 浏览器访问 http://127.0.0.1:5000 → '
    '分别选择正常、内圈故障、滚动体故障和外圈故障四类数据 → 执行故障诊断和健康预测 → 验证诊断结果、'
    '图表渲染和告警触发是否正确。'
)
add_para(
    '测试结果表明：四类数据的诊断结果均与预期一致；内圈故障数据触发了重度告警和紧急停机建议；'
    '正常数据诊断为正常状态，置信度 99.8%；三个 ECharts 图表均正常渲染，无 JavaScript 报错；'
    '诊断记录和告警记录正确写入数据库并在前端表格中展示。系统具备完整的工程可用性。'
)

# ========== 第8章 总结与展望 ==========
add_heading('第8章 总结与展望', level=1)

add_heading('8.1 工作总结', level=2)
add_para(
    '本课程设计完成了工业设备轴承故障智能监测与诊断系统的全流程开发，主要成果包括：'
)
add_para(
    '（1）基于 CWRU 公开数据集，完成了从信号预处理、特征提取到故障分类的完整算法链路，'
    '提取 20 维特征（12 时域 + 8 频域），采用随机森林分类器实现四类故障自动识别。'
)
add_para(
    '（2）发现并正确处理了滑动窗口重叠导致的数据泄漏问题。通过对比随机划分（0.997）和按文件分组（0.633）'
    '两种评估方法，揭示了随机划分评估的严重虚高问题，为课程实验和工业部署提供了正确的评估方法论。'
)
add_para(
    '（3）实现了基于 RMS 健康序列和多项式回归的健康趋势预测与阈值预警功能，'
    '实测内圈故障健康度 35.3，正确触发重度告警。'
)
add_para(
    '（4）构建了四类故障知识库，结合峭度严重度分级，输出包含成因、处理建议和检查清单的可解释运维决策。'
)
add_para(
    '（5）开发了 Flask + ECharts + SQLite 的 B/S 架构 Web 系统，提供 8 个 REST API，'
    '实现三图联动可视化和数据持久化，API 测试 12/12 通过，浏览器端四类故障验证正确。'
)

add_heading('8.2 不足与展望', level=2)
add_para(
    '本系统仍存在以下不足：'
    '（1）数据规模有限，仅使用了 20 个 CWRU 数据文件，滚动体和外圈故障样本偏少，'
    '按文件分组的泛化准确率（0.633）仍有提升空间；'
    '（2）仅使用了驱动端单通道数据，未利用多通道融合信息；'
    '（3）健康预测采用统计外推方法，未引入物理机理模型；'
    '（4）当前为离线文件分析模式，尚未支持实时数据流处理。'
)
add_para(
    '未来的改进方向包括：补充更多工况和故障类型的数据以提升模型泛化能力；'
    '引入一维卷积神经网络（1D-CNN）或 Transformer 等深度学习模型实现端到端特征学习；'
    '扩展为多通道数据融合诊断；接入实时传感器数据流实现在线监测；'
    '容器化部署以支持多设备并发监测。'
)

# ========== 参考文献 ==========
add_heading('参考文献', level=1)
refs = [
    '[1] Smith W A, Randall R B. Rolling element bearing diagnostics using the Case Western Reserve University data: A benchmark study[J]. Mechanical Systems and Signal Processing, 2015, 64-65: 100-131.',
    '[2] 雷亚国, 贾峰, 孔德同, 等. 大数据下机械智能故障诊断的机遇与挑战[J]. 机械工程学报, 2018, 54(5): 94-104.',
    '[3] Breiman L. Random forests[J]. Machine Learning, 2001, 45(1): 5-32.',
    '[4] 丁康, 李巍华, 朱小勇. 齿轮及齿轮箱故障诊断实用技术[M]. 北京: 机械工业出版社, 2005.',
    '[5] CWRU Bearing Data Center. Case Western Reserve University[EB/OL]. https://engineering.case.edu/bearingdatacenter.',
    '[6] 张昊. 工业设备轴承故障智能监测与诊断系统[CP/OL]. GitHub, https://github.com/uygby/zh111, 2026.',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.first_line_indent = Pt(-24)
    run = p.add_run(ref)
    set_font(run, '宋体', 'Times New Roman', 10.5)

# ========== 保存 ==========
doc.save(OUTPUT)
print(f"报告已生成: {OUTPUT}")
print(f"文件大小: {os.path.getsize(OUTPUT)} bytes")
